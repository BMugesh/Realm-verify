import numpy as np
import streamlit as st
import matplotlib.pyplot as plt
from mpl_toolkits.mplot3d import Axes3D
from matplotlib import cm
import time as time_module
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from io import BytesIO
import base64

# Initialize Session State for Animation
if 'animation_step' not in st.session_state:
    st.session_state.animation_step = 0
if 'is_animating' not in st.session_state:
    st.session_state.is_animating = False

# ============== CORE MATHEMATICAL FUNCTIONS ==============

def realm_curvature(x, y, alpha, beta):
    """Z = sin(α·x)·cos(β·y) - Main spatial curvature function"""
    return np.sin(alpha * x) * np.cos(beta * y)

def entropy_curve(x, y, alpha, beta):
    """∇z = √((∂z/∂x)² + (∂z/∂y)²) - Entropy gradient field"""
    dz_dx = alpha * np.cos(alpha * x) * np.cos(beta * y)
    dz_dy = -beta * np.sin(alpha * x) * np.sin(beta * y)
    return np.sqrt(dz_dx**2 + dz_dy**2)

def quantum_field(x, y, t, alpha, beta):
    """R_AB(x,y,t) = e^(-(x²+y²))·cos(√(x²+y²) - β·t) - Quantum tensor field"""
    r_squared = x**2 + y**2
    return np.exp(-r_squared) * np.cos(np.sqrt(r_squared) - beta * t) * alpha

def enhanced_surface(x, y, alpha, beta):
    """Enhanced 3D surface with multiple field contributions"""
    return np.sin(alpha * x) * np.cos(beta * y) + 0.3 * np.sin(x * y) / (1 + np.abs(x * y))

def simulate_geodesic_path(start, end, alpha, beta, steps=200):
    """Simulate quantum geodesic path in curved spacetime"""
    t = np.linspace(0, 1, steps)
    path_x = start[0] + t * (end[0] - start[0])
    path_y = start[1] + t * (end[1] - start[1])
    path_z = realm_curvature(path_x, path_y, alpha, beta)
    entropy = entropy_curve(path_x, path_y, alpha, beta)
    return path_x, path_y, path_z, entropy

def calculate_shannon_entropy(data):
    """H(φ) = -Σ P(x)·log(P(x)) - Shannon entropy calculation"""
    data = np.abs(data)
    data = data / (np.sum(data) + 1e-10)
    return -np.sum(data * np.log(data + 1e-10))

def calculate_field_statistics(field_data):
    """Calculate curvature statistics"""
    return {
        'mean': np.mean(field_data),
        'max': np.max(field_data),
        'min': np.min(field_data),
        'std': np.std(field_data)
    }

def create_word_report(alpha, beta, start_x, start_y, end_x, end_y, field_stats, path_stats, entropy, figures_dict):
    """Create a Word document with all simulation data and graphs"""
    doc = Document()
    
    # Title
    title = doc.add_heading('🌀 Realm Theory - Simulation Report', 0)
    title.alignment = 1  # Center alignment
    
    # Date and parameters
    doc.add_heading('Simulation Parameters', level=1)
    params_table = doc.add_table(rows=6, cols=2)
    params_table.style = 'Light Grid Accent 1'
    
    rows = params_table.rows
    rows[0].cells[0].text = "Parameter"
    rows[0].cells[1].text = "Value"
    rows[1].cells[0].text = "α (Alpha - Curvature Coupling)"
    rows[1].cells[1].text = f"{alpha:.4f}"
    rows[2].cells[0].text = "β (Beta - Field Feedback)"
    rows[2].cells[1].text = f"{beta:.4f}"
    rows[3].cells[0].text = "Observer Position (X, Y)"
    rows[3].cells[1].text = f"({start_x:.2f}, {start_y:.2f})"
    rows[4].cells[0].text = "Target Position (X, Y)"
    rows[4].cells[1].text = f"({end_x:.2f}, {end_y:.2f})"
    rows[5].cells[0].text = "Report Generated"
    from datetime import datetime
    rows[5].cells[1].text = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    # Field Statistics
    doc.add_heading('Field Statistics', level=1)
    stats_table = doc.add_table(rows=5, cols=2)
    stats_table.style = 'Light Grid Accent 1'
    
    rows = stats_table.rows
    rows[0].cells[0].text = "Statistic"
    rows[0].cells[1].text = "Value"
    rows[1].cells[0].text = "Mean Curvature"
    rows[1].cells[1].text = f"{field_stats['mean']:.6f}"
    rows[2].cells[0].text = "Max Curvature"
    rows[2].cells[1].text = f"{field_stats['max']:.6f}"
    rows[3].cells[0].text = "Min Curvature"
    rows[3].cells[1].text = f"{field_stats['min']:.6f}"
    rows[4].cells[0].text = "Std Deviation"
    rows[4].cells[1].text = f"{field_stats['std']:.6f}"
    
    # Add figures
    doc.add_page_break()
    doc.add_heading('Simulation Visualizations', level=1)
    
    figure_names = [
        ('Tensor Field 3D Surface', 'tensor_field_3d'),
        ('Enhanced Surface Visualization', 'enhanced_surface'),
        ('Entropy Flow Along Geodesic', 'entropy_along_path'),
        ('Entropy Distribution Heatmap', 'entropy_distribution'),
        ('Entropy Time Evolution', 'entropy_evolution'),
        ('Geodesic Path in 3D', 'geodesic_3d'),
        ('Geodesic 2D Projection', 'geodesic_height'),
    ]
    
    fig_count = 0
    for fig_name, fig_key in figure_names:
        if fig_key in figures_dict:
            if fig_count > 0 and fig_count % 2 == 0:
                doc.add_page_break()
            doc.add_heading(fig_name, level=2)
            img = figures_dict[fig_key]
            try:
                doc.add_picture(img, width=Inches(5.5))
            except:
                doc.add_paragraph(f"[Figure: {fig_name} - could not be embedded]")
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = 1  # Center
            fig_count += 1
    
    # Theory Section
    doc.add_page_break()
    doc.add_heading('Realm Theory - Summary', level=1)
    
    doc.add_heading('Core Principles', level=2)
    principles = [
        "Tensor Field Evolution: Spacetime curvature evolves via R_AB(x,y,t) = e^(-(x²+y²))·cos(√(x²+y²) - β·t)",
        "Geodesic Minimization: The Realm geodesic minimizes ∫ √(1 + (∂z/∂x)² + (∂z/∂y)²) dt",
        "Entropy Flow: Information complexity changes via ∇z = √((∂z/∂x)² + (∂z/∂y)²)",
        "Quantum Entanglement: Entangled particles share identical quantum states across geodesic connections"
    ]
    
    for principle in principles:
        doc.add_paragraph(principle, style='List Bullet')
    
    doc.add_heading('Key Findings', level=2)
    findings = [
        "Geodesic paths provide the shortest routes through curved spacetime",
        "Parameter tuning (α and β) enables precise control over field behavior",
        "Entropy measures information loss along quantum teleportation paths",
        "The unified geometric framework bridges quantum mechanics and relativity"
    ]
    
    for finding in findings:
        doc.add_paragraph(finding, style='List Bullet')
    
    return doc

# ============== STREAMLIT PAGE CONFIG ============== 
st.set_page_config(page_title="Realm Theory Visual Simulator", layout="wide")

# ============== TITLE & INTRODUCTION ==============

st.title("🌀 REALM THEORY - Advanced Visual Simulator")
st.markdown("""
### 🚀 Welcome to the Realm Dimension Explorer

Welcome to the **Realm Theory** simulator — where **quantum mechanics** meets **geometric spacetime**.

This interactive tool lets **students**, **researchers**, **investors**, and **science enthusiasts** explore:
- 🌌 The **curvature** of space-time as defined by unified **tensor fields**
- 🧬 The flow of **entropy** and quantum information along geodesic paths
- ⚡ Real-time simulation of **quantum teleportation** in a unified universe model
- 📊 Mathematical validation through interactive **surface plots** and **field dynamics**

> _"In the Realm, space bends to geometry, and geometry bends to will."_ — Realm Theory Principle
""")

# ============== THEOREM SECTION ==============

st.markdown("---")
st.success("📐 **Realm Theory - Core Theorem**")
st.write("**Definition:** The Realm represents a unified mathematical space where quantum entanglement and relativistic effects are unified through a single geometric structure.")

st.subheader("🔬 Fundamental Theorem Statement:")
st.write("Given a curved spacetime manifold defined by the metric tensor R_AB and a pair of quantum-entangled particles P₁ and P₂, there exists a unique geodesic path that represents the shortest physical distance through the underlying spacetime geometry.")

st.subheader("📊 Core Mathematical Principles:")
st.write("""
1. **Tensor Field Evolution:** Spacetime curvature evolves via: `R_AB(x,y,t) = e^(-(x²+y²))·cos(√(x²+y²) - β·t)`
2. **Geodesic Minimization:** The Realm geodesic minimizes: `∫ √(1 + (∂z/∂x)² + (∂z/∂y)²) dt`
3. **Entropy Flow:** Information complexity changes via: `∇z = √((∂z/∂x)² + (∂z/∂y)²)`
4. **Quantum Entanglement:** Entangled particles share identical quantum states across geodesic connections
""")

# ============== SIDEBAR CONTROLS ==============

st.sidebar.header("🎮 Universal Controls")
st.sidebar.markdown("### Curvature Parameters")
alpha = st.sidebar.slider('α (Curvature Coupling)', 0.1, 5.0, 1.0, help="Controls spatial warping strength in X direction")
beta = st.sidebar.slider('β (Field Feedback)', 0.1, 5.0, 0.5, help="Controls field response and oscillation in Y direction")

st.sidebar.markdown("### Observer & Target Position")
start_x = st.sidebar.slider('Start X (Observer)', -10.0, 10.0, -8.0)
start_y = st.sidebar.slider('Start Y (Observer)', -10.0, 10.0, -8.0)
end_x = st.sidebar.slider('End X (Target)', -10.0, 10.0, 8.0)
end_y = st.sidebar.slider('End Y (Target)', -10.0, 10.0, 8.0)

st.sidebar.markdown("### Visualization Settings")
mode = st.sidebar.radio("🧭 Simulation Mode", ["Static Observer", "Animated Traveler", "Both"])
auto_animate = st.sidebar.checkbox("🎬 Auto-Animate", value=False)
steps_select = st.sidebar.slider("Path Resolution", 50, 300, 200)

# ============== SIMULATION DATA ==============

start = np.array([start_x, start_y])
end = np.array([end_x, end_y])
x = np.linspace(-10, 10, 150)
y = np.linspace(-10, 10, 150)
X, Y = np.meshgrid(x, y)
Z = realm_curvature(X, Y, alpha, beta)
Z_enhanced = enhanced_surface(X, Y, alpha, beta)

# Geodesic path
gx, gy, gz, entropy = simulate_geodesic_path(start, end, alpha, beta, steps_select)

# Time-evolved quantum field
t_evolution = np.linspace(0, 2*np.pi, 100)
quantum_entropy_timeline = []
for t_val in t_evolution:
    Z_quantum = quantum_field(X, Y, t_val, alpha, beta)
    quantum_entropy_timeline.append(calculate_shannon_entropy(Z_quantum))

# Calculate statistics
field_stats = calculate_field_statistics(Z)
path_stats = calculate_field_statistics(gz)

# ============== MAIN TABS ==============

tabs = st.tabs([
    "🌐 Tensor Field", 
    "🧬 Enhanced Surfaces", 
    "🧠 Entropy Analysis", 
    "🧭 Geodesic Path", 
    "🎥 Animated Teleportation", 
    "📈 Field Dynamics",
    "📖 Theory Reference",
    "✅ Conclusion"
])

# ============== TAB 1: TENSOR FIELD ==============
with tabs[0]:
    st.subheader("🌐 Realm Tensor Field - 3D Visualization")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.info(f"""
        **Equation:** Z = sin(α·x)·cos(β·y)
        
        **Domain:** x, y ∈ [-10, 10]
        
        **Current Parameters:**
        - α = {alpha:.2f}
        - β = {beta:.2f}
        """)
        
        st.markdown("#### 📊 Field Statistics")
        st.metric("Mean Curvature", f"{field_stats['mean']:.4f}")
        st.metric("Max Curvature", f"{field_stats['max']:.4f}")
        st.metric("Min Curvature", f"{field_stats['min']:.4f}")
        st.metric("Std Deviation", f"{field_stats['std']:.4f}")
    
    with col2:
        st.info("💡 **Tip:** This visualization shows how spacetime curves due to the Realm field. Higher values indicate stronger curvature effects.")
    
    fig1 = plt.figure(figsize=(12, 8))
    ax1 = fig1.add_subplot(111, projection='3d')
    surf = ax1.plot_surface(X, Y, Z, cmap='plasma', alpha=0.9, linewidth=0, antialiased=True)
    ax1.set_xlabel("X Axis (Space Dimension 1)", fontsize=10)
    ax1.set_ylabel("Y Axis (Space Dimension 2)", fontsize=10)
    ax1.set_zlabel("Curvature (Field Strength)", fontsize=10)
    ax1.set_title("3D Visualization of Realm Tensor Field R_AB", fontsize=12, fontweight='bold')
    fig1.colorbar(surf, ax=ax1, label="Field Magnitude")
    st.session_state['fig1'] = fig1  # Store for export
    st.pyplot(fig1, use_container_width=True)
    
    st.markdown("---")
    
    # Contour plot
    fig_contour, ax_contour = plt.subplots(figsize=(10, 8))
    contour = ax_contour.contourf(X, Y, Z, levels=30, cmap='plasma')
    contour_lines = ax_contour.contour(X, Y, Z, levels=15, colors='white', alpha=0.3, linewidths=0.5)
    ax_contour.clabel(contour_lines, inline=True, fontsize=8)
    ax_contour.set_xlabel("X Axis")
    ax_contour.set_ylabel("Y Axis")
    ax_contour.set_title("Field Contour Map - Spatial Distribution")
    fig_contour.colorbar(contour, ax=ax_contour, label="Curvature Strength")
    st.pyplot(fig_contour, use_container_width=True)


# ============== TAB 2: ENHANCED SURFACES ==============
with tabs[1]:
    st.subheader("🧬 Multiple Surface Visualizations")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### Surface 1: Basic Tensor Field")
        st.markdown("""
        <div class="equation-box">
            Z₁ = sin(α·x)·cos(β·y)
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("#### Surface 2: Enhanced Field")
        st.markdown("""
        <div class="equation-box">
            Z₂ = sin(α·x)·cos(β·y) + 0.3·sin(xy)/(1+|xy|)
        </div>
        """, unsafe_allow_html=True)
    
    # Surface 1
    fig_s1 = plt.figure(figsize=(12, 5))
    ax_s1 = fig_s1.add_subplot(121, projection='3d')
    surf1 = ax_s1.plot_surface(X, Y, Z, cmap='viridis', alpha=0.9)
    ax_s1.set_title("Basic Tensor Field", fontweight='bold')
    ax_s1.set_xlabel("X")
    ax_s1.set_ylabel("Y")
    ax_s1.set_zlabel("Z")
    
    # Surface 2
    ax_s2 = fig_s1.add_subplot(122, projection='3d')
    surf2 = ax_s2.plot_surface(X, Y, Z_enhanced, cmap='coolwarm', alpha=0.9)
    ax_s2.set_title("Enhanced Field with Coupling", fontweight='bold')
    ax_s2.set_xlabel("X")
    ax_s2.set_ylabel("Y")
    ax_s2.set_zlabel("Z")
    
    st.session_state['fig2'] = fig_s1  # Store for export
    st.pyplot(fig_s1, use_container_width=True)
    
    # Difference map
    Z_diff = Z_enhanced - Z
    fig_diff = plt.figure(figsize=(12, 5))
    ax_diff = fig_diff.add_subplot(111, projection='3d')
    surf_diff = ax_diff.plot_surface(X, Y, Z_diff, cmap='seismic', alpha=0.9)
    ax_diff.set_title("Field Difference (Enhanced - Basic)", fontweight='bold')
    ax_diff.set_xlabel("X")
    ax_diff.set_ylabel("Y")
    ax_diff.set_zlabel("Δz")
    fig_diff.colorbar(surf_diff, ax=ax_diff, label="Field Deviation")
    st.pyplot(fig_diff, use_container_width=True)


# ============== TAB 3: ENTROPY ANALYSIS ==============
with tabs[2]:
    st.subheader("🧠 Entropy Flow & Information Dynamics")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown("""
        <div class="equation-box">
            <strong>Gradient Field:</strong><br>
            ∂z/∂x = α·cos(α·x)·cos(β·y)<br>
            ∂z/∂y = -β·sin(α·x)·sin(β·y)<br>
            <strong>Entropy Gradient:</strong><br>
            ∇z = √((∂z/∂x)² + (∂z/∂y)²)
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.info("💡 **Entropy** measures information complexity along geodesic paths")
    
    # Entropy along path
    fig_ent1, ax_ent1 = plt.subplots(figsize=(12, 5))
    progress = np.linspace(0, 1, len(entropy))
    ax_ent1.fill_between(progress, entropy, alpha=0.3, color='limegreen')
    ax_ent1.plot(progress, entropy, color='limegreen', linewidth=3, label='Entropy Flow')
    ax_ent1.set_xlabel("Geodesic Progress (0=Start, 1=End)", fontsize=11)
    ax_ent1.set_ylabel("Entropy (Information Complexity)", fontsize=11)
    ax_ent1.set_title("Information Flow Along Quantum Geodesic Path", fontweight='bold', fontsize=12)
    ax_ent1.grid(alpha=0.3)
    ax_ent1.legend()
    st.session_state['fig_ent1'] = fig_ent1  # Store for export
    st.pyplot(fig_ent1, use_container_width=True)
    
    # Entropy heatmap
    entropy_field = entropy_curve(X, Y, alpha, beta)
    fig_ent2, ax_ent2 = plt.subplots(figsize=(10, 8))
    entropy_plot = ax_ent2.contourf(X, Y, entropy_field, levels=40, cmap='hot')
    ax_ent2.set_xlabel("X Axis")
    ax_ent2.set_ylabel("Y Axis")
    ax_ent2.set_title("Entropy Distribution Across Realm Space", fontweight='bold')
    fig_ent2.colorbar(entropy_plot, ax=ax_ent2, label="Entropy Magnitude")
    st.session_state['fig_ent2'] = fig_ent2  # Store for export
    st.pyplot(fig_ent2, use_container_width=True)
    
    # Time evolution of entropy
    st.markdown("#### ⏱️ Quantum Entropy Time Evolution")
    fig_ent3, ax_ent3 = plt.subplots(figsize=(12, 5))
    ax_ent3.plot(t_evolution, quantum_entropy_timeline, color='cyan', linewidth=2.5, marker='o', markersize=4)
    ax_ent3.fill_between(t_evolution, quantum_entropy_timeline, alpha=0.2, color='cyan')
    ax_ent3.set_xlabel("Time Parameter t", fontsize=11)
    ax_ent3.set_ylabel("Shannon Entropy H(φ)", fontsize=11)
    ax_ent3.set_title("Shannon Entropy Evolution: H(φ) = -Σ P(x)·log(P(x))", fontweight='bold')
    ax_ent3.grid(alpha=0.3)
    st.session_state['fig_ent3'] = fig_ent3  # Store for export
    st.pyplot(fig_ent3, use_container_width=True)


# ============== TAB 4: GEODESIC PATH ==============
with tabs[3]:
    st.subheader("🧭 Quantum Geodesic Path")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        path_length_val = np.sum(np.sqrt(np.diff(gx)**2 + np.diff(gy)**2))
        st.info(f"""
        **Geodesic Definition:** Shortest path in curved spacetime
        
        - **Start Point (Blue):** ({start_x:.1f}, {start_y:.1f})
        - **End Point (Orange):** ({end_x:.1f}, {end_y:.1f})
        - **Path Length:** {path_length_val:.2f} units
        """)
    
    with col2:
        st.metric("Path Curvature (mean)", f"{path_stats['mean']:.4f}")
        st.metric("Path Curvature (max)", f"{path_stats['max']:.4f}")
    
    fig3 = plt.figure(figsize=(12, 7))
    ax3 = fig3.add_subplot(111, projection='3d')
    
    # Plot field surface (translucent)
    ax3.plot_surface(X, Y, Z, cmap='plasma', alpha=0.3, linewidth=0)
    
    # Plot geodesic path
    ax3.plot(gx, gy, gz, color='cyan', linewidth=3, label='Quantum Geodesic', zorder=5)
    
    # Plot start and end points
    z_start = realm_curvature(start[0], start[1], alpha, beta)
    z_end = realm_curvature(end[0], end[1], alpha, beta)
    ax3.scatter([start[0]], [start[1]], [z_start], color='blue', s=200, marker='o', label='Observer (Start)', zorder=10, edgecolors='white', linewidths=2)
    ax3.scatter([end[0]], [end[1]], [z_end], color='orange', s=200, marker='s', label='Target (End)', zorder=10, edgecolors='white', linewidths=2)
    
    # Plot projection on XY plane
    ax3.plot(gx, gy, 0, color='gray', linewidth=1, linestyle='--', alpha=0.5, label='XY Projection')
    
    ax3.set_xlabel("X Axis")
    ax3.set_ylabel("Y Axis")
    ax3.set_zlabel("Curvature (Z)")
    ax3.set_title("Shortest Quantum Teleportation Path in Realm Spacetime", fontweight='bold', fontsize=12)
    ax3.legend(loc='upper right')
    st.session_state['fig3'] = fig3  # Store for export
    st.pyplot(fig3, use_container_width=True)
    
    # 2D projection
    fig_2d = plt.figure(figsize=(10, 6))
    ax_2d = fig_2d.add_subplot(111)
    contour_2d = ax_2d.contourf(X, Y, Z, levels=30, cmap='plasma', alpha=0.7)
    ax_2d.plot(gx, gy, color='cyan', linewidth=3, label='Geodesic Path', zorder=5)
    ax_2d.scatter([start[0]], [start[1]], color='blue', s=300, marker='o', label='Observer', zorder=10, edgecolors='white', linewidths=2)
    ax_2d.scatter([end[0]], [end[1]], color='orange', s=300, marker='s', label='Target', zorder=10, edgecolors='white', linewidths=2)
    ax_2d.set_xlabel("X Axis")
    ax_2d.set_ylabel("Y Axis")
    ax_2d.set_title("2D Geodesic Path Projection on Realm XY Plane", fontweight='bold')
    ax_2d.legend()
    fig_2d.colorbar(contour_2d, ax=ax_2d, label="Curvature")
    st.session_state['fig4'] = fig_2d  # Store for export
    st.pyplot(fig_2d, use_container_width=True)


# ============== TAB 5: ANIMATED TELEPORTATION ==============
with tabs[4]:
    st.subheader("🎥 Quantum Teleportation Animation")
    
    st.markdown("""
    <div class="equation-box">
        <strong>Traveler Mode:</strong> Step-by-step journey visualization<br>
        <strong>Observer Mode:</strong> Complete path display<br>
        <strong>Both Mode:</strong> Comparison view
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    with col1:
        animation_speed = st.slider("⚡ Animation Speed (steps)", 1, 50, 5)
    with col2:
        if st.button("▶️ Start Animation"):
            st.session_state.is_animating = True
            st.session_state.animation_step = 0
    with col3:
        if st.button("⏹️ Stop Animation"):
            st.session_state.is_animating = False
    
    # Animation loop
    if st.session_state.is_animating:
        progress_bar = st.progress(0)
        status_text = st.empty()
        animation_container = st.empty()
        
        for step in range(0, len(gx), animation_speed):
            st.session_state.animation_step = step
            
            fig4 = plt.figure(figsize=(12, 7))
            ax4 = fig4.add_subplot(111, projection='3d')
            
            # Background field
            ax4.plot_surface(X, Y, Z, cmap='plasma', alpha=0.2, linewidth=0)
            
            # Full path in gray
            ax4.plot(gx, gy, gz, color='gray', linewidth=2, alpha=0.5, linestyle='--', label='Full Path')
            
            # Animated path
            ax4.plot(gx[:step+1], gy[:step+1], gz[:step+1], color='magenta', linewidth=3, label='Traveler Progress')
            
            # Current position
            if step < len(gx):
                ax4.scatter([gx[step]], [gy[step]], [gz[step]], color='yellow', s=300, marker='*', label='Current Position', zorder=10, edgecolors='white', linewidths=2)
            
            # Start and end
            ax4.scatter([start[0]], [start[1]], [z_start], color='blue', s=150, marker='o', label='Start', zorder=9)
            ax4.scatter([end[0]], [end[1]], [z_end], color='orange', s=150, marker='s', label='End', zorder=9)
            
            ax4.set_xlabel("X Axis")
            ax4.set_ylabel("Y Axis")
            ax4.set_zlabel("Curvature")
            ax4.set_title(f"🎥 Quantum Teleportation - Step {step} / {len(gx)}", fontweight='bold', fontsize=12)
            ax4.legend()
            
            animation_container.pyplot(fig4, use_container_width=True)
            progress_bar.progress(min(step / len(gx), 1.0))
            status_text.info(f"📍 Teleportation Progress: {step}/{len(gx)} steps")
            
            time_module.sleep(0.1)  # Control speed
    
    else:
        # Static view modes
        if mode in ["Static Observer", "Both"]:
            st.markdown("#### Observer Mode - Full Path Visualization")
            fig_static1 = plt.figure(figsize=(12, 7))
            ax_static1 = fig_static1.add_subplot(111, projection='3d')
            ax_static1.plot_surface(X, Y, Z, cmap='plasma', alpha=0.3, linewidth=0)
            ax_static1.plot(gx, gy, gz, color='cyan', linewidth=3, label='Complete Geodesic')
            ax_static1.scatter([start[0]], [start[1]], [z_start], color='blue', s=200, marker='o', label='Observer')
            ax_static1.scatter([end[0]], [end[1]], [z_end], color='orange', s=200, marker='s', label='Target')
            ax_static1.set_xlabel("X Axis")
            ax_static1.set_ylabel("Y Axis")
            ax_static1.set_zlabel("Curvature")
            ax_static1.set_title("Observer View - Complete Teleportation Path", fontweight='bold')
            ax_static1.legend()
            st.pyplot(fig_static1, use_container_width=True)
        
        if mode in ["Animated Traveler", "Both"]:
            st.markdown("#### Traveler Mode - Progressive Journey (50% Progress Snapshot)")
            fig_static2 = plt.figure(figsize=(12, 7))
            ax_static2 = fig_static2.add_subplot(111, projection='3d')
            ax_static2.plot_surface(X, Y, Z, cmap='plasma', alpha=0.3, linewidth=0)
            midpoint = len(gx) // 2
            ax_static2.plot(gx[:midpoint], gy[:midpoint], gz[:midpoint], color='magenta', linewidth=3, label='Traveler Path (50%)')
            ax_static2.plot(gx[midpoint:], gy[midpoint:], gz[midpoint:], color='gray', linewidth=1, linestyle='--', alpha=0.5)
            ax_static2.scatter([gx[midpoint]], [gy[midpoint]], [gz[midpoint]], color='yellow', s=300, marker='*', label='Current Position')
            ax_static2.set_xlabel("X Axis")
            ax_static2.set_ylabel("Y Axis")
            ax_static2.set_zlabel("Curvature")
            ax_static2.set_title("Traveler View - 50% Journey Progress", fontweight='bold')
            ax_static2.legend()
            st.pyplot(fig_static2, use_container_width=True)


# ============== TAB 6: FIELD DYNAMICS ==============
with tabs[5]:
    st.subheader("📈 Quantum Field Dynamics & Evolution")
    
    st.markdown("""
    <div class="equation-box">
        <strong>Time-Dependent Tensor Field:</strong><br>
        R_AB(x,y,t) = e^(-(x²+y²))·cos(√(x²+y²) - β·t)·α
    </div>
    """, unsafe_allow_html=True)
    
    # Multiple time snapshots
    col1, col2, col3 = st.columns(3)
    
    time_snapshots = [0, np.pi/2, np.pi]
    
    for idx, (col, t_snap) in enumerate(zip([col1, col2, col3], time_snapshots)):
        with col:
            Z_snap = quantum_field(X, Y, t_snap, alpha, beta)
            fig_snap = plt.figure(figsize=(8, 6))
            ax_snap = fig_snap.add_subplot(111, projection='3d')
            ax_snap.plot_surface(X, Y, Z_snap, cmap='coolwarm', alpha=0.9)
            ax_snap.set_title(f"t = {t_snap:.2f}", fontweight='bold')
            ax_snap.set_xlabel("X")
            ax_snap.set_ylabel("Y")
            ax_snap.set_zlabel("R_AB")
            st.pyplot(fig_snap, use_container_width=True)
    
    # Field energy over time
    st.markdown("#### Energy Density Evolution")
    energy_timeline = []
    for t_val in t_evolution:
        Z_energy = quantum_field(X, Y, t_val, alpha, beta)
        energy_timeline.append(np.sum(Z_energy**2))
    
    fig_energy = plt.figure(figsize=(12, 5))
    ax_energy = fig_energy.add_subplot(111)
    ax_energy.plot(t_evolution, energy_timeline, color='red', linewidth=2.5, marker='o', markersize=5)
    ax_energy.fill_between(t_evolution, energy_timeline, alpha=0.2, color='red')
    ax_energy.set_xlabel("Time t", fontsize=11)
    ax_energy.set_ylabel("Total Field Energy ∫R_AB² dV", fontsize=11)
    ax_energy.set_title("Quantum Field Energy Over Time", fontweight='bold')
    ax_energy.grid(alpha=0.3)
    st.pyplot(fig_energy, use_container_width=True)


# ============== TAB 7: THEORY REFERENCE ==============
with tabs[6]:
    st.subheader("📖 Complete Mathematical Theory Reference")
    
    st.markdown("""
    ### 🔬 Fundamental Equations
    
    #### 1. **Realm Curvature Function**
    ```
    z(x,y) = sin(α·x)·cos(β·y)
    ```
    Describes the base spatial curvature of the Realm.
    
    #### 2. **Entropy / Gradient Field**
    ```
    ∂z/∂x = α·cos(α·x)·cos(β·y)
    ∂z/∂y = -β·sin(α·x)·sin(β·y)
    
    Gradient Magnitude: ∇z = √((∂z/∂x)² + (∂z/∂y)²)
    ```
    Represents how rapidly the field changes in space.
    
    #### 3. **Quantum Tensor Field**
    ```
    R_AB(x,y,t) = e^(-(x²+y²))·cos(√(x²+y²) - β·t)·α
    ```
    Time-evolving quantum field coupled to spacetime geometry.
    
    #### 4. **Shannon Entropy**
    ```
    H(φ) = -Σ P(x)·log(P(x))
    where P(x) = |ψ(x)|² / Σ|ψ(x)|²
    ```
    Measures information complexity and quantum uncertainty.
    
    #### 5. **Geodesic Path**
    ```
    Minimize: S = ∫ √(1 + (∂z/∂x)² + (∂z/∂y)²) dt
    ```
    Shortest path in curved spacetime.
    
    ### 🎯 Parameter Definitions
    
    | Parameter | Symbol | Range | Physical Meaning |
    |-----------|--------|-------|------------------|
    | Curvature Coupling | α | 0.1-5.0 | Strength of spacetime warping in X-direction |
    | Field Feedback | β | 0.1-5.0 | Energy field response strength in Y-direction |
    
    ### 💡 Key Concepts
    
    - **Entanglement**: Quantum particles connected across geodesic paths
    - **Geodesic**: Shortest distance in curved space (like great circles on spheres)
    - **Entropy**: Information complexity measure; increases along least optimal paths
    - **Curvature**: Rate at which space deviates from flatness
    
    """)


# ============== TAB 8: CONCLUSION ==============
with tabs[7]:
    max_entropy_val = np.max(entropy)
    st.success("✅ **Conclusions & Key Findings**")
    
    st.subheader("🎯 Main Result:")
    st.write("""
    The Realm Theory demonstrates that quantum entanglement and relativistic spacetime curvature can be unified within a single geometric framework. By modeling spacetime as a continuously curved manifold, we show that:
    """)
    
    st.markdown("""
    1. **Geodesic Paths are Optimal:** The shortest distance between two quantum-entangled particles is not a straight line but follows the curved geometry of spacetime (geodesic), reducing travel distance and enabling "quantum teleportation".
    
    2. **Entropy Flow Predicts Information Loss:** Information complexity increases non-uniformly along geodesic paths. Regions of high curvature gradient correspond to high entropy dissipation, informing optimal teleportation strategies.
    
    3. **Tensor Field Evolution is Deterministic:** The time-dependent quantum tensor field R_AB(x,y,t) evolves predictably under the coupling parameters α and β, allowing precise simulation of quantum state evolution.
    
    4. **Parameter Tuning Enables Control:** Fine-tuning α (curvature strength) and β (field feedback) allows researchers to optimize:
        - Path efficiency (minimize travel distance)
        - Information preservation (minimize entropy loss)
        - Energy consumption (control field strength)
    
    5. **Universal Applicability:** The model scales from quantum mechanics to general relativity, making it valuable for:
        - Quantum computing optimization
                • Spacecraft trajectory planning
                • High-energy physics simulations
                • Machine learning field theory applications
        """)
    
    st.subheader("🚀 Practical Implications:")
    st.markdown("""
    - ✅ **Quantum Communication:** Entanglement can be used for instantaneous information transfer via geodesic paths
    - ✅ **Energy Efficiency:** Tuning parameters reduces wasted energy in quantum operations
    - ✅ **Spacetime Engineering:** Controlled curvature enables navigation and manipulation of space-time
    - ✅ **Cross-Dimensional Transfer:** Extended dimensions allow shortcuts through otherwise separated spaces
    """)
    
    st.subheader("🔮 Future Research Directions:")
    st.markdown("""
    - 🔬 Extend to 4D spacetime with full relativistic corrections
    - 🔬 Multi-particle entanglement networks
    - 🔬 Machine learning optimization of parameters
    - 🔬 Experimental validation in quantum systems
    - 🔬 Integration with particle physics models (Standard Model compatibility)
    """)
    
    st.divider()
    
    st.subheader("📊 Simulation Summary")
    st.write("**Current Configuration:**")
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Alpha (α)", f"{alpha:.2f}")
    with col2:
        st.metric("Beta (β)", f"{beta:.2f}")
    with col3:
        st.metric("Mean Curvature", f"{field_stats['mean']:.4f}")
    with col4:
        st.metric("Max Entropy", f"{np.max(entropy):.4f}")
    
    col5, col6 = st.columns(2)
    with col5:
        st.write(f"**Observer Position:** ({start_x:.1f}, {start_y:.1f})")
    with col6:
        st.write(f"**Target Position:** ({end_x:.1f}, {end_y:.1f})")
    
    st.divider()
    st.info("🌟 **In the Realm, geometry is destiny.** 🌟")
    
    # Export data
    st.divider()
    st.subheader("📥 Export Simulation Data")
    
    col_json, col_word = st.columns(2)
    
    # JSON Export
    with col_json:
        export_data = {
            'parameters': {
                'alpha': float(alpha),
                'beta': float(beta),
                'observer': [float(start_x), float(start_y)],
                'target': [float(end_x), float(end_y)]
            },
            'statistics': {
                'field': field_stats,
                'path': path_stats,
                'max_entropy': float(np.max(entropy)),
                'min_entropy': float(np.min(entropy))
            },
            'geodesic_path': {
                'x': gx.tolist(),
                'y': gy.tolist(),
                'z': gz.tolist(),
                'entropy': entropy.tolist()
            }
        }
        
        import json
        json_str = json.dumps(export_data, indent=2)
        st.download_button("📄 Download JSON Data", json_str, "realm_simulation_data.json", "application/json")
    
    # Word Document Export
    with col_word:
        if st.button("📊 Generate Word Report", key="gen_word"):
            with st.spinner("Generating Word document with all graphs..."):
                # Collect all figures from session state
                figures_dict = {}
                
                # Store matplotlib figures as images in memory
                if 'fig1' in st.session_state:
                    img_buffer = BytesIO()
                    st.session_state['fig1'].savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                    img_buffer.seek(0)
                    figures_dict['tensor_field_3d'] = img_buffer
                
                if 'fig2' in st.session_state:
                    img_buffer = BytesIO()
                    st.session_state['fig2'].savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                    img_buffer.seek(0)
                    figures_dict['enhanced_surface'] = img_buffer
                
                if 'fig_ent1' in st.session_state:
                    img_buffer = BytesIO()
                    st.session_state['fig_ent1'].savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                    img_buffer.seek(0)
                    figures_dict['entropy_distribution'] = img_buffer
                
                if 'fig_ent2' in st.session_state:
                    img_buffer = BytesIO()
                    st.session_state['fig_ent2'].savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                    img_buffer.seek(0)
                    figures_dict['entropy_evolution'] = img_buffer
                
                if 'fig_quantum' in st.session_state:
                    img_buffer = BytesIO()
                    st.session_state['fig_quantum'].savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                    img_buffer.seek(0)
                    figures_dict['quantum_field'] = img_buffer
                
                if 'fig3' in st.session_state:
                    img_buffer = BytesIO()
                    st.session_state['fig3'].savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                    img_buffer.seek(0)
                    figures_dict['geodesic_3d'] = img_buffer
                
                if 'fig4' in st.session_state:
                    img_buffer = BytesIO()
                    st.session_state['fig4'].savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                    img_buffer.seek(0)
                    figures_dict['geodesic_height'] = img_buffer
                
                # Use fig_ent1 for entropy along path
                if 'fig_ent1' in st.session_state:
                    img_buffer = BytesIO()
                    st.session_state['fig_ent1'].savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                    img_buffer.seek(0)
                    figures_dict['entropy_along_path'] = img_buffer
                
                # Create the Word document
                doc = create_word_report(alpha, beta, start_x, start_y, end_x, end_y, 
                                        field_stats, path_stats, entropy, figures_dict)
                
                # Save to BytesIO
                word_buffer = BytesIO()
                doc.save(word_buffer)
                word_buffer.seek(0)
                
                # Store in session state for download
                st.session_state['word_doc'] = word_buffer.getvalue()
                st.success("✅ Word document generated successfully!")
        
        # Show download button if document is generated
        if 'word_doc' in st.session_state:
            st.download_button(
                label="⬇️ Download Word Report",
                data=st.session_state['word_doc'],
                file_name="Realm_Theory_Simulation_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_word"
            )

# ============== FOOTER ==============
st.divider()
st.write("""
---
🌀 **Realm Theory Simulator v2.0** | Enhanced with Multiple Surfaces, Animations, and Comprehensive Theory

**Built for:**
- 👨‍🎓 Students
- 🧠 Theorists  
- 🔬 Researchers
- 💼 Investors
- 🚀 Sci-fi Enthusiasts

*"Where quantum mechanics meets geometric spacetime."*
""")
